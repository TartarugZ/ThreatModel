from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from functions.functions import beautify_ubi
from data_base import db_controller

document = Document()
section = document.sections[0]
# левое поле в миллиметрах
section.left_margin = Mm(30)
# правое поле в миллиметрах
section.right_margin = Mm(10)
# верхнее поле в миллиметрах
section.top_margin = Mm(20)
# нижнее поле в миллиметрах
section.bottom_margin = Mm(20)


def add_heading(name, size, level):
    heading = document.add_heading(name, level=level)
    heading.style.font.name = 'Times New Roman'
    heading.style.font.size = Pt(size)
    heading.style.font.color.rgb = RGBColor(0, 0, 0)
    return heading


def add_paragraph(text):
    paragraph = document.add_paragraph(text)
    paragraph.style.font.name = 'Times New Roman'
    paragraph.style.font.size = Pt(14)
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Mm(1)
    paragraph.paragraph_format.space_after = Mm(1)
    return paragraph


def add_from_file(name):
    with open(name, 'r', encoding='utf-8') as f:
        data = f.readlines()
        for i in data:
            if '\n' in i:
                i = i.replace('\n', '')
                add_paragraph(i)


def to_view(objects, namae=''):
    final_objects = []
    for j in objects:
        final_objects.append(str(j[0]))
    final_objects = list(set(final_objects))
    result = ''
    for j in final_objects:
        result += namae + j + ', '
    return result


def create_threat_model(engine):
    doc_name = add_heading('МОДЕЛЬ УГРОЗ БЕЗОПАСНОСТИ', 25, 0)
    doc_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # document.add_page_break()

    heading1_definitions = add_heading('1. ТЕРМИНЫ И ОПРЕДЕЛЕНИЯ', 16, 1)

    add_from_file('threat_model_prompt/1.txt')

    heading2_general = add_heading('2. ОБЩИЕ ПОЛОЖЕНИЯ', 16, 1)

    heading2_1 = add_heading('2.1. Введение', 14, 2)
    add_from_file('threat_model_prompt/2_1.txt')

    heading2_2 = add_heading('2.2. Оцениваемые угрозы', 14, 2)
    add_from_file('threat_model_prompt/2_2.txt')

    heading2_3 = add_heading('2.3. Ответственность за обеспечение защиты информации (безопасности)', 14, 2)
    add_from_file('threat_model_prompt/2_3.txt')
    table_responsibility = document.add_table(rows=2, cols=3)
    heading_cells = table_responsibility.rows[0].cells
    heading_cells[0].text = '№\nп/п'
    heading_cells[1].text = 'Роль подразделения /\nдолжностного лица'
    heading_cells[2].text = 'Должностное лицо / подразделение'
    prompt_cells = table_responsibility.rows[1].cells
    prompt_cells[0].text = '1'
    prompt_cells[1].text = 'Пример: Ответственный за обеспечение безопасности'
    prompt_cells[2].text = 'Пример: Руководитель отдела средств защиты'
    table_responsibility.style = 'Table Grid'
    for cell in table_responsibility.columns[0].cells:
        cell.width = Cm(2)
    for cell in table_responsibility.columns[1].cells:
        cell.width = Cm(8)
    for cell in table_responsibility.columns[2].cells:
        cell.width = Cm(8)
    for row in table_responsibility.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    heading2_4 = add_heading('2.4. Особенности пересмотра Модели угроз', 14, 2)
    add_from_file('threat_model_prompt/2_4.txt')

    heading3_objects = add_heading('3. ОПИСАНИЕ СИСТЕМ И СЕТЕЙ И ИХ ХАРАКТЕРИСТИКА КАК ОБЪЕКТОВ ЗАЩИТЫ', 16, 1)

    heading3_1 = add_heading('3.1 Общее описание объекта оценки угроз', 14, 2)
    add_from_file('threat_model_prompt/3_1.txt')

    heading3_2 = add_heading('3.2. Состав и архитектура объекта оценки', 14, 2)
    add_from_file('threat_model_prompt/3_2_1.txt')
    table_consist = document.add_table(rows=5, cols=3)
    heading_cells = table_consist.rows[0].cells
    heading_cells[0].text = '№\nп/п'
    heading_cells[1].text = 'Характеристика'
    heading_cells[2].text = 'Значение характеристики'
    prompt_cells = table_consist.rows[1].cells
    prompt_cells[0].text = '1'
    prompt_cells[1].text = 'Программно-аппаратные средства'
    devices = db_controller.get_all_devices(engine=engine)
    temp = []
    dictionary = {}
    for i in devices:
        temp.append(i.description)
    temp_set = temp.copy()
    temp_set = list(set(temp_set))
    for i in temp_set:
        count = temp.count(i)
        dictionary[i] = count
    result = ''
    for i in dictionary.items():
        result += f'{i[0]} – {i[1]}\n'
    prompt_cells[2].text = result
    prompt_cells = table_consist.rows[2].cells
    prompt_cells[0].text = '2'
    prompt_cells[1].text = 'Общесистемное программное обеспечение'
    prompt_cells[
        2].text = 'Пример:\nОперационные системы:\n- Debian GNU/Linux;\n- Microsoft Windows Server 2019 Standart, русская версия;\n- Microsoft Windows Server 2012 R2 Standart x64;\n- Microsoft Windows 10 Pro, 64-разрядная'
    prompt_cells = table_consist.rows[3].cells
    prompt_cells[0].text = '3'
    prompt_cells[1].text = 'Прикладное программное обеспечение'
    prompt_cells[2].text = 'Пример:\n- 1С Бухгалтерия'
    prompt_cells = table_consist.rows[4].cells
    prompt_cells[0].text = '4'
    prompt_cells[1].text = 'Средства защиты информации'
    prompt_cells[
        2].text = 'Пример:\nСредства антивирусной защиты: Kaspersky Endpoint Security для Windows (версия 11.1.1.126)\nСредства криптографической защиты информации: Программный комплекс ViPNet Client 4 (версия 4.5)'
    table_consist.style = 'Table Grid'
    for cell in table_consist.columns[0].cells:
        cell.width = Cm(2)
    for cell in table_consist.columns[1].cells:
        cell.width = Cm(8)
    for cell in table_consist.columns[2].cells:
        cell.width = Cm(8)
    for row in table_consist.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_from_file('threat_model_prompt/3_2_2.txt')
    table_network = document.add_table(rows=2, cols=5)
    heading_cells = table_network.rows[0].cells
    heading_cells[0].text = '№\nп/п'
    heading_cells[1].text = 'Категория сети электросвязи'
    heading_cells[2].text = 'Наименование оператора связи'
    heading_cells[3].text = 'Цель взаимодействия с сетью электросвязи'
    heading_cells[4].text = 'Способ взаимодействия с сетью электросвязи'
    prompt_cells = table_network.rows[1].cells
    prompt_cells[0].text = '1'
    prompt_cells[1].text = 'Общего пользования'
    prompt_cells[2].text = 'ПАО КОМПАНИЯ'
    prompt_cells[3].text = 'Оказание услуг'
    prompt_cells[4].text = 'Тип доступа: проводной, беспроводной. Протоколы: HTTP, POP3, FTP, SMTP, IMAP4, TCP/IP'
    table_network.style = 'Table Grid'
    for cell in table_network.columns[0].cells:
        cell.width = Cm(2)
    for cell in table_network.columns[1].cells:
        cell.width = Cm(5)
    for cell in table_network.columns[2].cells:
        cell.width = Cm(4)
    for cell in table_network.columns[3].cells:
        cell.width = Cm(4)
    for cell in table_network.columns[4].cells:
        cell.width = Cm(5)
    for row in table_network.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_from_file('threat_model_prompt/3_2_3.txt')
    table_technologies = document.add_table(rows=17, cols=3)
    heading_cells = table_technologies.rows[0].cells
    heading_cells[0].text = '№\nп/п'
    heading_cells[1].text = 'Технология'
    heading_cells[2].text = 'Использование'
    prompt_cells = table_technologies.rows[1].cells
    prompt_cells[0].text = '1'
    prompt_cells[1].text = 'Съемные носители информации'
    prompt_cells[2].text = 'Не используются'
    prompt_cells = table_technologies.rows[2].cells
    prompt_cells[0].text = '2'
    prompt_cells[1].text = 'Технология виртуализации'
    prompt_cells[2].text = 'Используются'
    prompt_cells = table_technologies.rows[3].cells
    prompt_cells[0].text = '3'
    prompt_cells[1].text = 'Технология беспроводного доступа'
    prompt_cells = table_technologies.rows[4].cells
    prompt_cells[0].text = '4'
    prompt_cells[1].text = 'Мобильные технические средства'
    prompt_cells = table_technologies.rows[5].cells
    prompt_cells[0].text = '5'
    prompt_cells[1].text = 'Веб-серверы'
    prompt_cells = table_technologies.rows[6].cells
    prompt_cells[0].text = '6'
    prompt_cells[1].text = 'Технология веб-доступа'
    prompt_cells = table_technologies.rows[7].cells
    prompt_cells[0].text = '7'
    prompt_cells[1].text = 'Smart-карты'
    prompt_cells = table_technologies.rows[8].cells
    prompt_cells[0].text = '8'
    prompt_cells[1].text = 'Технологии грид-систем'
    prompt_cells = table_technologies.rows[9].cells
    prompt_cells[0].text = '9'
    prompt_cells[1].text = 'Технологии суперкомпьютерных систем'
    prompt_cells = table_technologies.rows[10].cells
    prompt_cells[0].text = '10'
    prompt_cells[1].text = 'Большие данные'
    prompt_cells = table_technologies.rows[11].cells
    prompt_cells[0].text = '11'
    prompt_cells[1].text = 'Числовое программное оборудование'
    prompt_cells = table_technologies.rows[12].cells
    prompt_cells[0].text = '12'
    prompt_cells[1].text = 'Электронная почта'
    prompt_cells = table_technologies.rows[13].cells
    prompt_cells[0].text = '13'
    prompt_cells[1].text = 'Технология передачи видеоинформации'
    prompt_cells = table_technologies.rows[14].cells
    prompt_cells[0].text = '14'
    prompt_cells[1].text = 'Технология удаленного доступа'
    prompt_cells = table_technologies.rows[15].cells
    prompt_cells[0].text = '15'
    prompt_cells[1].text = 'Технология удаленного внеполосного доступа'
    prompt_cells = table_technologies.rows[16].cells
    prompt_cells[0].text = '16'
    prompt_cells[1].text = 'Технология искусственного интеллекта'
    table_technologies.style = 'Table Grid'
    for cell in table_technologies.columns[0].cells:
        cell.width = Cm(2)
    for cell in table_technologies.columns[1].cells:
        cell.width = Cm(10)
    for cell in table_technologies.columns[2].cells:
        cell.width = Cm(6)
    for row in table_technologies.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    heading4_negative = add_heading(
        '4. ВОЗМОЖНЫЕ НЕГАТИВНЫЕ ПОСЛЕДСТВИЯ ОТ РЕАЛИЗАЦИИ (ВОЗНИКНОВЕНИЯ) УГРОЗ БЕЗОПАСНОСТИ ИНФОРМАЦИИ', 16, 1)
    add_from_file('threat_model_prompt/4.txt')
    negatives = db_controller.get_all_negatives(engine)
    damage_types = {0: 'Ущерб физическому лицу', 1: 'Ущерб юридическому лицу', 2: 'Ущерб государству'}
    table_negatives = document.add_table(rows=len(negatives) + 1, cols=3)
    heading_cells = table_negatives.rows[0].cells
    heading_cells[0].text = 'Иденти-\nфикатор'
    heading_cells[1].text = 'Негативные последствия'
    heading_cells[2].text = 'Вид риска (ущерба)'
    for i in range(len(negatives)):
        neg_cells = table_negatives.rows[i + 1].cells
        neg_cells[0].text = f'НП.{i + 1}'
        neg_cells[1].text = f'{negatives[i].name}'
        neg_cells[2].text = f'{damage_types[int(negatives[i].damage_type)]}'
    table_negatives.style = 'Table Grid'
    for cell in table_negatives.columns[0].cells:
        cell.width = Cm(3)
    for cell in table_negatives.columns[1].cells:
        cell.width = Cm(10)
    for cell in table_negatives.columns[2].cells:
        cell.width = Cm(6)
    for row in table_negatives.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    heading5_objects = add_heading(
        '5. ВОЗМОЖНЫЕ ОБЪЕКТЫ ВОЗДЕЙСТВИЯ УГРОЗ БЕЗОПАСНОСТИ ИНФОРМАЦИИ', 16, 1)
    add_from_file('threat_model_prompt/5_1.txt')
    impact_types = db_controller.get_all_impact_types(engine)
    table_impact = document.add_table(rows=len(impact_types) + 1, cols=2)
    heading_cells = table_impact.rows[0].cells
    heading_cells[0].text = 'Иденти-\nфикатор'
    heading_cells[1].text = 'Вид воздействия'
    for i in range(len(impact_types)):
        neg_cells = table_impact.rows[i + 1].cells
        neg_cells[0].text = f'ВВ.{i + 1}'
        neg_cells[1].text = f'{impact_types[i].name}'
    table_impact.style = 'Table Grid'
    for cell in table_impact.columns[0].cells:
        cell.width = Cm(5)
    for cell in table_impact.columns[1].cells:
        cell.width = Cm(20)
    for row in table_impact.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_from_file('threat_model_prompt/5_2.txt')
    neg_obj_imp = db_controller.get_negative_object_impact(engine)
    set_neg = []
    for i in neg_obj_imp:
        set_neg.append(i[0])
    set_neg = list(set(set_neg))
    sort_objects = {}
    for i in set_neg:
        sort_objects[i] = ''
    neg_obj_imp = db_controller.get_negative_object_impact(engine)
    for i in neg_obj_imp:
        sort_objects[i[0]] = sort_objects[i[0]] + i[1] + '|'
    set_obj = []
    for i in sort_objects.items():
        temp = i[1].split('|')
        temp = set(temp)
        for j in temp:
            set_obj.append(j)
    set_obj = set(set_obj)
    sorting_impacts = {}
    set_obj.remove('')
    for i in set_obj:
        sorting_impacts[i] = ''
    neg_obj_imp = db_controller.get_negative_object_impact(engine)
    for i in neg_obj_imp:
        sorting_impacts[i[1]] = sorting_impacts[i[1]] + str(i[2]) + '|'
    for i in sorting_impacts.items():
        temp = i[1].split('|')
        temp.remove('')
        for k in range(len(temp)):
            temp[k] = int(temp[k])
        temp.sort()
        temp = set(temp)
        sorting_impacts[i[0]] = ''
        for j in temp:
            sorting_impacts[i[0]] = sorting_impacts[i[0]] + 'ВВ.' + str(j) + ' '
    for i in sort_objects.items():
        all_obj = i[1].split('|')
        all_obj.remove('')
        all_obj = set(all_obj)
        sort_objects[i[0]] = ''
        for m in all_obj:
            sort_objects[i[0]] = sort_objects[i[0]] + m + '|'
    lines = []
    intervals = []
    for i in range(len(set_neg)):
        obj = sort_objects[set_neg[i]]
        obj = obj.split('|')
        obj.remove('')
        used = False
        for j in obj:
            impacts = sorting_impacts[j]
            if used:
                lines.append(('', j, impacts))
            else:
                lines.append((set_neg[i], j, impacts))
                used = True
                intervals.append(len(lines))
    table_negative_device_type_impact = document.add_table(rows=len(lines) + 1, cols=3)
    heading_cells = table_negative_device_type_impact.rows[0].cells
    heading_cells[0].text = 'Негативные последствия'
    heading_cells[1].text = 'Объекты воздействия'
    heading_cells[2].text = 'Виды воздействия'
    for i in range(len(lines)):
        neg_cells = table_negative_device_type_impact.rows[i + 1].cells
        neg_cells[0].text = f'{lines[i][0]}'
        neg_cells[1].text = f'{lines[i][1]}'
        neg_cells[2].text = f'{lines[i][2]}'
    table_negative_device_type_impact.style = 'Table Grid'
    for row in table_negative_device_type_impact.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i in range(len(intervals)):
        if i != 0:
            table_negative_device_type_impact.cell(intervals[i - 1], 0).merge(
                table_negative_device_type_impact.cell(intervals[i] - 1, 0))
    table_negative_device_type_impact.cell(intervals[len(intervals) - 1], 0).merge(
        table_negative_device_type_impact.cell(len(lines), 0))

    heading6_threat_source = add_heading(
        '6. ИСТОЧНИКИ УГРОЗ БЕЗОПАСНОСТИ ИНФОРМАЦИИ', 16, 1)
    heading6_1 = add_heading(
        '6.1. Антропогенные источники', 14, 2)
    add_from_file('threat_model_prompt/6_1_1.txt')
    table_intruders = document.add_table(rows=14, cols=4)
    heading_cells = table_intruders.rows[0].cells
    heading_cells[0].text = '№\nп/п'
    heading_cells[1].text = 'Вид нарушителя'
    heading_cells[2].text = 'Возможные цели реализации\nугроз безопасности информации'
    heading_cells[3].text = 'Актуальность\nвозможных нарушителей'
    intruders = [
        'Специальные службы иностранных государств',
        'Террористические, экстремистские группировки',
        'Преступные группы (криминальные структуры)',
        'Отдельные физические лица (хакеры)',
        'Конкурирующие организации',
        'Разработчики программных, программно-аппаратных средств',
        'Лица, обеспечивающие поставку программных, программно-аппаратных средств, обеспечивающих систем',
        'Поставщики вычислительных услуг, услуг связи',
        'Лица, привлекаемые для установки, настройки, испытаний, пусконаладочных и иных видов работ',
        'Лица, обеспечивающие функционирование систем и сетей или обеспечивающие системы оператора',
        'Авторизованные пользователи систем и сетей',
        'Системные администраторы и администраторы безопасности',
        'Бывшие (уволенные) работники (пользователи)',
    ]
    aims = [
        'Нанесение ущерба государству в области обороны, безопасности и правопорядка, а также в иных отдельных областях его деятельности или секторах экономики; Дискредитация деятельности отдельных органов государственной власти, организаций; Получение конкурентных преимуществ на уровне государства; Срыв заключения международных договоров; Создание внутриполитического кризиса',
        'Совершение террористических актов, угроза жизни граждан; Нанесение ущерба отдельным сферам деятельности или секторам экономики государства; Дестабилизация общества; Дестабилизация деятельности органов государственной власти, организаций',
        'Получение финансовой или иной материальной выгоды; Желание самореализации (подтверждение статуса)',
        'Получение финансовой или иной материальной выгоды;Любопытство или желание самореализации (подтверждение статуса)',
        'Получение финансовой или иной материальной выгоды; Получение конкурентных преимуществ',
        'Получение финансовой или иной материальной выгоды; Получение конкурентных преимуществ; Внедрение дополнительных функциональных возможностей в программные или программно-аппаратные средства на этапе разработки; Непреднамеренные, неосторожные или неквалифицированные действия',
        'Получение финансовой или иной материальной выгоды; Получение конкурентных преимуществ; Непреднамеренные, неосторожные или неквалифицированные действия',
        'Получение финансовой или иной материальной выгоды; Получение конкурентных преимуществ; Непреднамеренные, неосторожные или неквалифицированные действия',
        'Получение финансовой или иной материальной выгоды; Получение конкурентных преимуществ; Непреднамеренные, неосторожные или неквалифицированные действия',
        'Получение финансовой или иной материальной выгоды; Непреднамеренные, неосторожные или неквалифицированные действия',
        'Получение финансовой или иной материальной выгоды; Любопытство или желание самореализации (подтверждение статуса); Непреднамеренные, неосторожные или неквалифицированные действия; Месть за ранее совершенные действия',
        'Получение финансовой или иной материальной выгоды; Любопытство или желание самореализации (подтверждение статуса); Непреднамеренные, неосторожные или неквалифицированные действия; Месть за ранее совершенные действия',
        'Получение финансовой или иной материальной выгоды; Месть за ранее совершенные действия',
    ]
    for i in range(len(intruders)):
        intruder_cells = table_intruders.rows[i + 1].cells
        intruder_cells[0].text = f'{i + 1}'
        intruder_cells[1].text = f'{intruders[i]}'
        intruder_cells[2].text = f'{aims[i]}'
        intruder_cells[3].text = f'Актуально/\nНе актуально (причина)'
    table_intruders.style = 'Table Grid'
    for cell in table_intruders.columns[0].cells:
        cell.width = Cm(1)
    for cell in table_intruders.columns[1].cells:
        cell.width = Cm(5)
    for cell in table_intruders.columns[2].cells:
        cell.width = Cm(15)
    for cell in table_intruders.columns[3].cells:
        cell.width = Cm(4)
    for row in table_intruders.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_from_file('threat_model_prompt/6_1_2.txt')
    table_actual_intruders = document.add_table(rows=3, cols=4)
    heading_cells = table_actual_intruders.rows[0].cells
    heading_cells[0].text = '№\nп/п'
    heading_cells[1].text = 'Вид нарушителя'
    heading_cells[2].text = 'Категория'
    heading_cells[3].text = 'Уровень возможностей'
    intruder_cells = table_actual_intruders.rows[1].cells
    intruder_cells[0].text = f'1'
    intruder_cells[1].text = f'Отдельные физические лица (хакеры)'
    intruder_cells[2].text = f'Внешний'
    intruder_cells[3].text = f'H1. Нарушитель, обладающий базовыми возможностями'
    intruder_cells = table_actual_intruders.rows[2].cells
    intruder_cells[0].text = f'2'
    intruder_cells[1].text = f'Системные администраторы и администраторы безопасности'
    intruder_cells[2].text = f'Внутренний'
    intruder_cells[3].text = f'Н2. Нарушитель, обладающий базовыми повышенными возможностями'
    table_actual_intruders.style = 'Table Grid'
    for cell in table_actual_intruders.columns[0].cells:
        cell.width = Cm(1)
    for cell in table_actual_intruders.columns[1].cells:
        cell.width = Cm(10)
    for cell in table_actual_intruders.columns[2].cells:
        cell.width = Cm(5)
    for cell in table_actual_intruders.columns[3].cells:
        cell.width = Cm(8)
    for row in table_actual_intruders.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_from_file('threat_model_prompt/6_1_3.txt')

    heading7_negative = add_heading(
        '7. СПОСОБЫ РЕАЛИЗАЦИИ (ВОЗНИКНОВЕНИЯ) УГРОЗ БЕЗОПАСНОСТИ ИНФОРМАЦИИ', 16, 1)
    add_from_file('threat_model_prompt/7_1.txt')
    realization_ways = db_controller.get_all_realization_ways(engine)
    table_realization_ways = document.add_table(rows=len(realization_ways) + 1, cols=2)
    heading_cells = table_realization_ways.rows[0].cells
    heading_cells[0].text = 'Идентификатор'
    heading_cells[1].text = 'Способы реализации'
    for i in range(len(realization_ways)):
        realization_cells = table_realization_ways.rows[i + 1].cells
        realization_cells[0].text = f'СР.{i + 1}'
        realization_cells[1].text = f'{realization_ways[i].name}'
    table_realization_ways.style = 'Table Grid'
    for cell in table_realization_ways.columns[0].cells:
        cell.width = Cm(5)
    for cell in table_realization_ways.columns[1].cells:
        cell.width = Cm(20)
    for row in table_realization_ways.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_from_file('threat_model_prompt/7_2.txt')
    obj_realization = db_controller.get_object_realization(engine)
    set_obj = []
    for i in obj_realization:
        set_obj.append(i[0])
    set_obj = list(set(set_obj))
    sort_obj = {}
    for i in set_obj:
        sort_obj[i] = ''
    obj_realization = db_controller.get_object_realization(engine)
    for i in obj_realization:
        sort_obj[i[0]] = sort_obj[i[0]] + str(i[1]) + '|'
    for i in sort_obj.items():
        all_realization = i[1].split('|')
        all_realization.remove('')
        all_realization = set(all_realization)
        sort_obj[i[0]] = ''
        for m in all_realization:
            sort_obj[i[0]] = sort_obj[i[0]] + 'СР.' + m + ' '
    table_obj_realization = document.add_table(rows=len(set_obj) + 1, cols=2)
    heading_cells = table_obj_realization.rows[0].cells
    heading_cells[0].text = 'Объект взаимодействия'
    heading_cells[1].text = 'Способы реализации'
    for i in range(len(set_obj)):
        realization_cells = table_obj_realization.rows[i + 1].cells
        realization_cells[0].text = f'{set_obj[i]}'
        realization_cells[1].text = f'{sort_obj[set_obj[i]]}'
    table_obj_realization.style = 'Table Grid'
    for row in table_obj_realization.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_paragraph(
        'Соотношение актуальных объектов реализации, выявленных у них интерфейсов и способов реализации угроз безопасности информации представлено в таблице 12.')
    add_paragraph('Таблица 12. Возможные способы реализации угроз безопасности на интерфейсах объектов взаимодействия')
    obj_interface_realization = db_controller.get_object_interface_realization(engine)
    set_obj = []
    for i in obj_interface_realization:
        set_obj.append(i[0])
    set_obj = list(set(set_obj))
    sort_interface = {}
    for i in set_obj:
        sort_interface[i] = ''
    obj_interface_realization = db_controller.get_object_interface_realization(engine)
    for i in obj_interface_realization:
        sort_interface[i[0]] = sort_interface[i[0]] + i[1] + '|'
    set_interface = []
    for i in sort_interface.items():
        temp = i[1].split('|')
        temp = set(temp)
        for j in temp:
            set_interface.append(j)
    set_interface = set(set_interface)
    sorting_realization = {}
    set_interface.remove('')
    for i in set_interface:
        sorting_realization[i] = ''
    obj_interface_realization = db_controller.get_object_interface_realization(engine)
    for i in obj_interface_realization:
        sorting_realization[i[1]] = sorting_realization[i[1]] + str(i[2]) + '|'
    for i in sorting_realization.items():
        temp = i[1].split('|')
        temp.remove('')
        for k in range(len(temp)):
            temp[k] = int(temp[k])
        temp.sort()
        temp = set(temp)
        sorting_realization[i[0]] = ''
        for j in temp:
            sorting_realization[i[0]] = sorting_realization[i[0]] + 'СР.' + str(j) + ' '
    for i in sort_interface.items():
        all_obj = i[1].split('|')
        all_obj.remove('')
        all_obj = set(all_obj)
        sort_interface[i[0]] = ''
        for m in all_obj:
            sort_interface[i[0]] = sort_interface[i[0]] + m + '|'
    lines = []
    intervals = []
    for i in range(len(set_obj)):
        obj = sort_interface[set_obj[i]]
        obj = obj.split('|')
        obj.remove('')
        used = False
        for j in obj:
            realization = sorting_realization[j]
            if used:
                lines.append(('', j, realization))
            else:
                lines.append((set_obj[i], j, realization))
                used = True
                intervals.append(len(lines))
    table_obj_interface_realization = document.add_table(rows=len(lines) + 1, cols=3)
    heading_cells = table_obj_interface_realization.rows[0].cells
    heading_cells[0].text = 'Объекты взаимодействия'
    heading_cells[1].text = 'Интерфейсы'
    heading_cells[2].text = 'Способы реализации'
    for i in range(len(lines)):
        oir_cells = table_obj_interface_realization.rows[i + 1].cells
        oir_cells[0].text = f'{lines[i][0]}'
        oir_cells[1].text = f'{lines[i][1]}'
        oir_cells[2].text = f'{lines[i][2]}'
    table_obj_interface_realization.style = 'Table Grid'
    for row in table_obj_interface_realization.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i in range(len(intervals)):
        if i != 0:
            table_obj_interface_realization.cell(intervals[i - 1], 0).merge(
                table_obj_interface_realization.cell(intervals[i] - 1, 0))
    table_obj_interface_realization.cell(intervals[len(intervals) - 1], 0).merge(
        table_obj_interface_realization.cell(len(lines), 0))

    heading8_threats = add_heading(
        '8. АКТУАЛЬНЫЕ УГРОЗЫ БЕЗОПАСНОСТИ ИНФОРМАЦИИ', 16, 1)
    add_from_file('threat_model_prompt/8_1.txt')
    external_intruder = {0: 'Внешний нарушитель с базовыми возможностями',
                         1: 'Внешний нарушитель с базовыми повышенными возможностями',
                         2: 'Внешний нарушитель со средними возможностями',
                         3: 'Внешний нарушитель с высокими возможностями'}
    internal_intruder = {0: 'Внутренний нарушитель с базовыми возможностями',
                         1: 'Внутренний нарушитель с базовыми повышенными возможностями',
                         2: 'Внутренний нарушитель со средними возможностями',
                         3: 'Внутренний нарушитель с высокими возможностями'}
    used_ubi = db_controller.get_used_ubi(engine)
    list_used_ubi = []
    id_used_ubi = []
    for i in used_ubi:
        id_used_ubi.append(i[0])
        list_used_ubi.append(i)
    id_used_ubi = list(set(id_used_ubi))
    list_used_ubi = list(set(list_used_ubi))
    table = document.add_table(rows=len(list_used_ubi) + 1, cols=3)
    heading_cells = table.rows[0].cells
    heading_cells[0].text = 'Идентификатор\nугрозы'
    heading_cells[1].text = 'Наименование угрозы'
    heading_cells[2].text = 'Характеристика нарушителя,\nнеобходимая для реализации угрозы'
    heading_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_cells[2].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(len(list_used_ubi)):
        base_cells = table.rows[i + 1].cells
        base_cells[0].text = f'УБИ.{beautify_ubi(list_used_ubi[i][0])}'
        base_cells[1].text = f'{list_used_ubi[i][1]}'
        base_cells[2].text = f'{internal_intruder[list_used_ubi[i][2]]}\n\n{external_intruder[list_used_ubi[i][3]]}'
    table.style = 'Table Grid'
    add_paragraph('Соотношение угрозы, объектов, на которые она распространяется, негативных последствий ее реализации, а также способов ее реализации и сценариев представлено в таблице 14.')
    add_paragraph('Таблица 14. Результаты оценки возможных угроз безопасности информации')
    table = document.add_table(rows=len(list_used_ubi) + 1, cols=5)
    heading_cells = table.rows[0].cells
    heading_cells[0].text = 'Идентификатор\nугрозы'
    heading_cells[1].text = 'Объекты воздействия'
    heading_cells[2].text = 'Негативные последствия'
    heading_cells[3].text = 'Способы реализации угрозы'
    heading_cells[4].text = 'Возможные сценарии реализации угрозы'
    heading_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_cells[1].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_cells[2].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_cells[3].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_cells[4].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    id_used_ubi.sort()
    for i in range(len(id_used_ubi)):
        base_cells = table.rows[i + 1].cells
        base_cells[0].text = f'УБИ.{beautify_ubi(id_used_ubi[i])}'
        objects = db_controller.get_obj_by_ubi_id(id_used_ubi[i], engine=engine)
        result = to_view(objects)
        base_cells[1].text = f'{result[:-2]}'
        negatives = db_controller.get_negative_by_ubi_id(id_used_ubi[i], engine=engine)
        result = to_view(negatives, 'НП.')
        base_cells[2].text = f'{result[:-2]}'
        realization_ways = db_controller.get_realization_by_ubi_id(id_used_ubi[i], engine=engine)
        result = to_view(realization_ways,  'СР.')
        base_cells[3].text = f'{result[:-2]}'
        scenario = db_controller.get_scenario_by_ubi_id(id_used_ubi[i], engine=engine)
        scenarios = {}
        unique_scenario = []
        for j in scenario:
            unique_scenario.append(j[0])
        unique_scenario = list(set(unique_scenario))
        for j in unique_scenario:
            scenarios[j] = ''
        scenario = db_controller.get_scenario_by_ubi_id(id_used_ubi[i], engine=engine)
        for j in scenario:
            scenarios[j[0]] += '-T' + str(j[1]) + '.' + str(j[2]) + '\n'
        result = ''
        for j in scenarios.items():
            result += j[1] + '\n'
        base_cells[4].text = f'{result}'
    table.style = 'Table Grid'

    document.save('files/demo.docx')
